"""DOCX document parser."""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.core.exceptions import InvalidDocumentError
from app.core.logging import get_logger
from app.services.ingestion.base import DocumentParser, ExtractedSection, NormalizedDocument

logger = get_logger(__name__)


class DocxParser(DocumentParser):
    """Parse DOCX files into heading-aware sections."""

    @property
    def file_type(self) -> str:
        return "docx"

    def parse(self, content: bytes, *, filename: str) -> NormalizedDocument:
        if not content:
            raise InvalidDocumentError(
                "Document file is empty",
                details={"reason": "empty_file", "file_type": self.file_type},
            )

        try:
            document = DocxDocument(BytesIO(content))
        except PackageNotFoundError as exc:
            raise InvalidDocumentError(
                "DOCX file is corrupted or unreadable",
                details={"reason": "corrupted_file", "file_type": self.file_type},
            ) from exc
        except Exception as exc:
            raise InvalidDocumentError(
                "DOCX file is corrupted or unreadable",
                details={"reason": "corrupted_file", "file_type": self.file_type},
            ) from exc

        sections: list[ExtractedSection] = []
        current_heading: str | None = None
        current_lines: list[str] = []

        def flush_section() -> None:
            nonlocal current_lines, current_heading
            body = "\n".join(line for line in current_lines if line).strip()
            if not body:
                current_lines = []
                return
            sections.append(
                ExtractedSection(
                    text=body,
                    section_index=len(sections),
                    page_number=1,
                    section=current_heading,
                )
            )
            current_lines = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name if paragraph.style is not None else "").lower()
            if "heading" in style_name:
                flush_section()
                current_heading = text
                current_lines = [text]
            else:
                current_lines.append(text)

        flush_section()

        if not sections:
            raise InvalidDocumentError(
                "DOCX file contains no extractable content",
                details={"reason": "empty_text", "file_type": self.file_type},
            )

        logger.info(
            "docx_parsed",
            extra={
                "operation": "parse_docx",
                "source": filename,
                "section_count": len(sections),
            },
        )
        return NormalizedDocument(
            sections=sections,
            file_type=self.file_type,
            source=filename,
            page_count=1,
            section_count=len(sections),
        )
